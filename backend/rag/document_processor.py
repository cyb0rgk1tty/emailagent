"""
Document processing module for RAG system
Supports PDF, DOCX, Markdown, and TXT files
"""
import os
import re
from pathlib import Path
from typing import List, Dict, Optional
import logging

# Import document loaders based on availability
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process various document formats and extract text content"""

    SUPPORTED_EXTENSIONS = {
        '.md': 'markdown',
        '.txt': 'text',
        '.pdf': 'pdf',
        '.docx': 'docx'
    }

    def __init__(self, knowledge_base_path: str = "/app/knowledge"):
        """Initialize document processor

        Args:
            knowledge_base_path: Path to knowledge base directory
        """
        self.knowledge_base_path = Path(knowledge_base_path)
        logger.info(f"Initialized DocumentProcessor with path: {knowledge_base_path}")

    def is_supported(self, file_path: Path) -> bool:
        """Check if file format is supported

        Args:
            file_path: Path to file

        Returns:
            True if file format is supported
        """
        extension = file_path.suffix.lower()

        # Check if extension is supported
        if extension not in self.SUPPORTED_EXTENSIONS:
            return False

        # Check if required dependencies are available
        if extension == '.pdf' and not PDF_AVAILABLE:
            logger.warning("PDF support not available - install PyPDF2")
            return False

        if extension == '.docx' and not DOCX_AVAILABLE:
            logger.warning("DOCX support not available - install python-docx")
            return False

        return True

    def extract_text_from_markdown(self, file_path: Path) -> str:
        """Extract text from markdown file

        Args:
            file_path: Path to markdown file

        Returns:
            Extracted text content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # Markdown is already text-based, return as-is
            # Could add markdown parsing for structured extraction later
            return content

        except Exception as e:
            logger.error(f"Error reading markdown file {file_path}: {e}")
            return ""

    def extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from plain text file

        Args:
            file_path: Path to text file

        Returns:
            Extracted text content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return ""

    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content
        """
        if not PDF_AVAILABLE:
            logger.error("PyPDF2 not installed - cannot process PDF files")
            return ""

        try:
            text_parts = []

            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)

                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()

                    if text:
                        text_parts.append(text)

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {e}")
            return ""

    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not installed - cannot process DOCX files")
            return ""

        try:
            doc = docx.Document(file_path)

            # Extract text from paragraphs
            text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}")
            return ""

    def extract_text(self, file_path: Path) -> str:
        """Extract text from any supported document format

        Args:
            file_path: Path to document

        Returns:
            Extracted text content
        """
        extension = file_path.suffix.lower()

        if extension == '.md':
            return self.extract_text_from_markdown(file_path)
        elif extension == '.txt':
            return self.extract_text_from_txt(file_path)
        elif extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif extension == '.docx':
            return self.extract_text_from_docx(file_path)
        else:
            logger.warning(f"Unsupported file format: {extension}")
            return ""

    def get_document_type(self, file_path: Path) -> str:
        """Determine document type based on directory structure

        Args:
            file_path: Path to document

        Returns:
            Document type (product_catalog, pricing, certification, capability, faq)
        """
        # Get relative path from knowledge base root
        try:
            rel_path = file_path.relative_to(self.knowledge_base_path)
            parts = rel_path.parts

            if len(parts) > 1:
                # First directory is the document type
                category = parts[0].lower()

                # Map directory names to document types
                type_mapping = {
                    'products': 'product_catalog',
                    'pricing': 'pricing',
                    'certifications': 'certification',
                    'capabilities': 'capability',
                    'faq': 'faq'
                }

                return type_mapping.get(category, 'capability')

        except ValueError:
            # File is not under knowledge base path
            pass

        return 'capability'

    def extract_sections(self, text: str) -> List[Dict[str, str]]:
        """Extract sections from markdown-style text with headers

        Args:
            text: Text content with markdown headers

        Returns:
            List of sections with title and content
        """
        sections = []

        # Split by markdown headers (## or ###)
        header_pattern = r'^(#{1,3})\s+(.+)$'
        lines = text.split('\n')

        current_section = None
        current_content = []

        for line in lines:
            match = re.match(header_pattern, line)

            if match:
                # Save previous section
                if current_section:
                    sections.append({
                        'title': current_section,
                        'content': '\n'.join(current_content).strip()
                    })

                # Start new section
                current_section = match.group(2).strip()
                current_content = []
            else:
                current_content.append(line)

        # Add final section
        if current_section:
            sections.append({
                'title': current_section,
                'content': '\n'.join(current_content).strip()
            })

        # If no sections found, return entire text as one section
        if not sections:
            sections.append({
                'title': 'Document',
                'content': text.strip()
            })

        return sections

    def process_document(self, file_path: Path) -> Dict:
        """Process a document and extract metadata

        Args:
            file_path: Path to document

        Returns:
            Document data with metadata
        """
        if not self.is_supported(file_path):
            logger.warning(f"File not supported: {file_path}")
            return None

        # Extract text content
        text = self.extract_text(file_path)

        if not text:
            logger.warning(f"No text extracted from {file_path}")
            return None

        # Extract sections
        sections = self.extract_sections(text)

        # Determine document type
        doc_type = self.get_document_type(file_path)

        return {
            'file_path': str(file_path),
            'file_name': file_path.name,
            'document_type': doc_type,
            'full_text': text,
            'sections': sections,
            'metadata': {
                'file_size': file_path.stat().st_size,
                'file_extension': file_path.suffix,
                'section_count': len(sections)
            }
        }

    def scan_knowledge_base(self) -> List[Dict]:
        """Scan knowledge base directory for all supported documents

        Returns:
            List of processed documents
        """
        documents = []

        if not self.knowledge_base_path.exists():
            logger.error(f"Knowledge base path not found: {self.knowledge_base_path}")
            return documents

        # Recursively scan for supported files
        for file_path in self.knowledge_base_path.rglob('*'):
            if file_path.is_file() and self.is_supported(file_path):
                logger.info(f"Processing document: {file_path}")

                doc_data = self.process_document(file_path)

                if doc_data:
                    documents.append(doc_data)

        logger.info(f"Scanned knowledge base: {len(documents)} documents found")
        return documents


# Singleton instance
_processor = None

def get_document_processor() -> DocumentProcessor:
    """Get singleton document processor instance"""
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor
